#!/usr/bin/env python3
"""
Aplica el formato institucional IREM/BID a un .docx que ya existe.

Lee el documento de entrada, clasifica cada bloque por su papel (titulo,
seccion, subtitulo, cuerpo, vineta, tabla) y lo vuelve a escribir sobre
plantilla.docx. Asi el resultado hereda el encabezado con los logos, los
margenes, la tipografia y el espaciado del formato, sin arrastrar los estilos
del documento original.

No cambia el contenido: conserva el texto, las negritas y las cursivas tal
como estaban. Lo unico que cambia es el formato.

Uso:
    uv run --with python-docx python aplicar_formato.py ENTRADA.docx [SALIDA.docx] [--fuente Calibri]
"""
import re
import sys
from pathlib import Path

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn

sys.path.insert(0, str(Path(__file__).resolve().parent))
from irem_docx import Documento  # noqa: E402

# un subtitulo no deberia ser mas largo que esto; si lo es, es cuerpo en negrita
LARGO_MAX_TITULO = 120


def bloques(doc):
    for hijo in doc.element.body.iterchildren():
        if hijo.tag == qn("w:p"):
            yield Paragraph(hijo, doc)
        elif hijo.tag == qn("w:tbl"):
            yield Table(hijo, doc)


def marca(rpr, tag):
    if rpr is None:
        return False
    n = rpr.find(qn("w:" + tag))
    return n is not None and n.get(qn("w:val")) not in ("0", "false")


def runs_de(p):
    """
    [(texto, negrita, cursiva), ...] con las marcas heredadas del estilo.

    Recorre todos los w:r del parrafo, no solo los hijos directos, para no
    perder el texto que va dentro de hipervinculos, smartTags o controles de
    contenido. El texto borrado (w:delText) y los campos sin texto (w:pgNum,
    numeros de pagina sueltos) no aportan nada y quedan fuera solos.
    """
    est = p.style
    neg_est = bool(est is not None and est.font.bold)
    cur_est = bool(est is not None and est.font.italic)
    salida = []
    for r in p._p.iter(qn("w:r")):
        texto = "".join(t.text or "" for t in r.iter(qn("w:t")))
        if r.find(qn("w:tab")) is not None:
            texto += "\t"
        if not texto:
            continue
        rpr = r.find(qn("w:rPr"))
        salida.append((texto,
                       marca(rpr, "b") or neg_est,
                       marca(rpr, "i") or cur_est))
    return salida


def props(p):
    ppr = p._p.find(qn("w:pPr"))
    def g(tag, attr="val"):
        if ppr is None:
            return None
        n = ppr.find(qn("w:" + tag))
        return n.get(qn("w:" + attr)) if n is not None else None
    num = None
    if ppr is not None and ppr.find(qn("w:numPr")) is not None:
        n = ppr.find(qn("w:numPr")).find(qn("w:numId"))
        num = n.get(qn("w:val")) if n is not None else None
    return dict(estilo=(p.style.name if p.style is not None else ""),
                jc=g("jc"), num=num,
                salto=bool(p._p.findall(".//" + qn("w:br"))))


def nivel_de_estilo(nombre):
    """Devuelve el papel si el estilo es un titulo de Word, o None."""
    n = (nombre or "").lower()
    if n in ("title", "titulo", "titulo del documento"):
        return "titulo"
    m = re.match(r"(?:heading|titulo|t.tulo)\s*(\d)", n)
    if m:
        return {1: "h1", 2: "h2", 3: "h3"}.get(int(m.group(1)), "h3")
    return None


def clasifica(p, d, runs, primero):
    texto = "".join(t for t, _, _ in runs).strip()
    if not texto:
        return "salto" if d["salto"] else "vacio"

    por_estilo = nivel_de_estilo(d["estilo"])
    if por_estilo:
        return por_estilo

    negs = {b for _, b, _ in runs}
    curs = {i for _, _, i in runs}
    corto = len(texto) <= LARGO_MAX_TITULO

    if d["num"]:
        # una lista numerada en negrita y corta es una seccion; el resto, vinetas
        return "h1" if (negs == {True} and corto) else "vineta"
    if "list" in (d["estilo"] or "").lower() or "lista" in (d["estilo"] or "").lower():
        return "vineta"
    if primero and d["jc"] == "center" and negs == {True}:
        return "titulo"
    if negs == {True} and corto:
        return "h2"
    if curs == {True} and corto:
        return "h3"
    # "Etiqueta: valor" con la etiqueta en negrita
    if runs and runs[0][1] and not all(negs) and ":" in runs[0][0]:
        return "campo"
    return "cuerpo"


def celdas_de(tabla):
    filas = []
    for row in tabla.rows:
        fila = []
        for cell in row.cells:
            partes = []
            for p in cell.paragraphs:
                t = "".join(t for t, _, _ in runs_de(p)).strip()
                if t:
                    partes.append(t)
            fila.append("\n".join(partes))
        filas.append(fila)
    return filas


def es_fila_encabezado(tabla, i):
    """Gris si Word la marca como encabezado, o si esta sombreada, o si va toda en negrita."""
    row = tabla.rows[i]
    trpr = row._tr.find(qn("w:trPr"))
    if trpr is not None and trpr.find(qn("w:tblHeader")) is not None:
        return True
    sombreada = True
    negrita = True
    hay_texto = False
    for cell in row.cells:
        tcpr = cell._tc.find(qn("w:tcPr"))
        shd = tcpr.find(qn("w:shd")) if tcpr is not None else None
        relleno = shd.get(qn("w:fill")) if shd is not None else None
        if relleno in (None, "auto", "FFFFFF"):
            sombreada = False
        for p in cell.paragraphs:
            for t, b, _ in runs_de(p):
                if t.strip():
                    hay_texto = True
                    if not b:
                        negrita = False
    return hay_texto and (sombreada or negrita)


def anchos_de(tabla):
    grid = tabla._tbl.find(qn("w:tblGrid"))
    if grid is None:
        return None
    pesos = [int(c.get(qn("w:w")) or 0) for c in grid]
    return pesos if all(p > 0 for p in pesos) else None


def alineaciones_de(tabla):
    """Toma la alineacion de la primera fila de datos, columna por columna."""
    for i, row in enumerate(tabla.rows):
        if es_fila_encabezado(tabla, i):
            continue
        salida = []
        for cell in row.cells:
            jc = None
            for p in cell.paragraphs:
                ppr = p._p.find(qn("w:pPr"))
                n = ppr.find(qn("w:jc")) if ppr is not None else None
                if n is not None:
                    jc = n.get(qn("w:val"))
                    break
            salida.append(jc)
        return [j if j in ("both", "center", "left", "right") else
                ("both" if k == 0 else "center") for k, j in enumerate(salida)]
    return None


def reformatea(entrada, salida, fuente):
    origen = Document(str(entrada))
    doc = Documento(fuente=fuente)
    primero = True
    cuenta = {}
    for b in bloques(origen):
        if isinstance(b, Table):
            filas = celdas_de(b)
            if not any(any(c.strip() for c in f) for f in filas):
                continue
            encs = {i for i in range(len(b.rows)) if es_fila_encabezado(b, i)} or {0}
            doc.tabla(filas, encabezados=encs, pesos=anchos_de(b),
                      alineaciones=alineaciones_de(b))
            cuenta["tabla"] = cuenta.get("tabla", 0) + 1
            primero = False
            continue

        d = props(b)
        runs = runs_de(b)
        papel = clasifica(b, d, runs, primero)
        if papel == "vacio":
            continue          # el ritmo vertical lo repone el formato
        cuenta[papel] = cuenta.get(papel, 0) + 1
        if papel == "salto":
            doc.salto_pagina()
            continue
        if papel == "campo":
            corte = runs[0][0]
            resto = [(t, b_, i_) for t, b_, i_ in runs[1:]]
            doc.campo(corte, resto)
        else:
            getattr(doc, papel)(runs)
        if papel != "titulo":
            primero = False
    doc.guardar(salida)
    return salida, cuenta


def main(argv):
    args = list(argv[1:])
    fuente = "Calibri"
    if "--fuente" in args:
        k = args.index("--fuente")
        if k + 1 >= len(args):
            sys.exit("Falta el nombre de la fuente despues de --fuente")
        fuente = args[k + 1]
        del args[k:k + 2]
    if not args:
        print(__doc__)
        sys.exit(1)
    entrada = Path(args[0])
    salida = Path(args[1]) if len(args) > 1 else \
        entrada.with_name(entrada.stem + "-formateado.docx")
    if salida.resolve() == entrada.resolve():
        sys.exit("La salida no puede ser el archivo de entrada")
    salida, cuenta = reformatea(entrada, salida, fuente)
    resumen = ", ".join(f"{k}={v}" for k, v in sorted(cuenta.items()))
    print(f"Formateado con fuente '{fuente}': {salida}")
    print(f"Bloques: {resumen}")


if __name__ == "__main__":
    main(sys.argv)
